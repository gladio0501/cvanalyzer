"""
CV Document Parser Module.
"""

import pdfplumber
import docx

def parse_pdf(file_path):
    """
    Extract text content from PDF files using pdfplumber.
    """
    with pdfplumber.open(file_path) as pdf:
        return "\n".join(page.extract_text() or "" for page in pdf.pages)

def parse_docx(file_path):
    """
    Extract text content from Microsoft Word DOCX files.
    """
    doc = docx.Document(file_path)
    print(f"Document text: {[para.text for para in doc.paragraphs]}")
    return "\n".join([para.text for para in doc.paragraphs])




def parse_cv(file_path):
    """
    Main entry point for CV document parsing with format detection.
    """
    if file_path.endswith('.pdf'):
        return parse_pdf(file_path)
    elif file_path.endswith('.docx'):
        return parse_docx(file_path)
    else:
        raise ValueError("Unsupported file type")
